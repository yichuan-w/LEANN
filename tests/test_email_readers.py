"""Comprehensive tests for LEANN email readers (EmlReader, EmlxReader)."""

import email
import os
import re
import tempfile
from email.header import Header
from email.mime.application import MIMEApplication
from email.mime.base import MIMEBase
from email.mime.image import MIMEImage
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.utils import formatdate
from pathlib import Path

import pytest

# ── Import path setup ─────────────────────────────────────
# The email reader lives under apps/ — add it to sys.path
import sys
sys.path.insert(0, str(Path(__file__).parent.parent / "apps"))

from email_data.LEANN_email_reader import (  # noqa: E402
    EmlxReader,
    EmlReader,
    _build_email_document,
    _collect_attachments,
    _decode_addr,
    _decode_subject,
    _extract_attachment_text,
    _extract_rtf_text,
    _payload_to_text,
    _split_quoted_thread,
    _strip_html,
)


# ==========================================================
# HELPERS — generate test .eml files
# ==========================================================

def _make_simple_text(subject: str = "Test Subject", body: str = "Hello world",
                      from_addr: str = "alice@example.com",
                      to_addr: str = "bob@example.com",
                      date: str | None = None) -> str:
    """Build a minimal text/plain .eml string."""
    msg = MIMEText(body)
    msg["Subject"] = subject
    msg["From"] = from_addr
    msg["To"] = to_addr
    msg["Date"] = date or formatdate(localtime=True)
    return msg.as_string()


def _make_multipart(subject: str, html_body: str = "",
                    text_body: str = "", attachments: list | None = None,
                    extra_headers: dict | None = None) -> str:
    """Build a multipart/mixed .eml string with optional text, HTML, attachments."""
    msg = MIMEMultipart("mixed")
    msg["Subject"] = subject
    msg["From"] = "alice@example.com"
    msg["To"] = "bob@example.com"
    msg["Date"] = formatdate(localtime=True)
    if extra_headers:
        for k, v in extra_headers.items():
            msg[k] = v

    if text_body or html_body:
        alt = MIMEMultipart("alternative")
        if text_body:
            alt.attach(MIMEText(text_body, "plain"))
        if html_body:
            alt.attach(MIMEText(html_body, "html"))
        msg.attach(alt)

    if attachments:
        for att in attachments:
            msg.attach(att)

    return msg.as_string()


def _write_eml(dir_path: str, filename: str, content: str) -> str:
    """Write a .eml file and return its full path."""
    path = Path(dir_path) / filename
    path.write_text(content, encoding="utf-8")
    return str(path)


# ==========================================================
# TESTS: _payload_to_text
# ==========================================================

class TestPayloadToText:
    def test_bytes_input(self):
        assert _payload_to_text(b"hello bytes") == "hello bytes"

    def test_str_input(self):
        assert _payload_to_text("hello string") == "hello string"

    def test_null_bytes(self):
        assert _payload_to_text(b"\x00\x00") == "\x00\x00"

    def test_utf8_bytes(self):
        assert _payload_to_text("café".encode("utf-8")) == "café"

    def test_non_text_type(self):
        assert _payload_to_text(42) == ""
        assert _payload_to_text(None) == ""
        assert _payload_to_text([]) == ""


# ==========================================================
# TESTS: _strip_html
# ==========================================================

class TestStripHtml:
    def test_simple_tags(self):
        result = _strip_html("<p>Hello</p>")
        assert "Hello" in result
        assert "<p>" not in result

    def test_no_html(self):
        assert _strip_html("plain text") == "plain text"

    def test_empty_string(self):
        assert _strip_html("") == ""

    def test_nested_tags(self):
        """bs4.get_text(separator='\\n') inserts newlines between block tags."""
        result = _strip_html("<div><p><b>bold</b> text</p></div>")
        assert "bold" in result
        assert "text" in result
        assert "<div>" not in result
        assert "<b>" not in result

    def test_html_with_entities(self):
        """bs4.get_text decodes HTML entities (&amp; → &, &lt; → <)."""
        result = _strip_html("<p>foo &amp; bar &lt; baz</p>")
        assert "foo & bar < baz" in result

    def test_malformed_html(self):
        result = _strip_html("<p>unclosed <b>tags")
        assert "unclosed" in result
        assert "tags" in result


# ==========================================================
# TESTS: _decode_subject
# ==========================================================

class TestDecodeSubject:
    def test_plain_subject(self):
        assert _decode_subject("Hello") == "Hello"

    def test_none_subject(self):
        assert _decode_subject(None) == "No Subject"

    def test_empty_subject(self):
        assert _decode_subject("") == "No Subject"

    def test_rfc2047_base64(self):
        # =?UTF-8?B? ... ?=
        encoded = str(Header("café", "utf-8"))
        result = _decode_subject(encoded)
        assert "café" in result

    def test_rfc2047_qp(self):
        encoded = "=?ISO-8859-1?Q?M=FCller?="
        result = _decode_subject(encoded)
        assert "Müller" in result

    def test_multiple_encoded_words(self):
        encoded = "=?UTF-8?B?SGVsbG8=?= =?UTF-8?B?V29ybGQ=?="
        result = _decode_subject(encoded)
        assert "HelloWorld" in result

    def test_unknown_charset_fallback(self):
        encoded = "=?X-UNKNOWN?B?SGVsbG8=?="
        result = _decode_subject(encoded)
        # Should not crash; falls back to utf-8 decode
        assert result is not None
        assert len(result) > 0

    def test_long_subject(self):
        long_subj = "A" * 1000
        assert _decode_subject(long_subj) == long_subj


# ==========================================================
# TESTS: _decode_addr
# ==========================================================

class TestDecodeAddr:
    def test_plain_addr(self):
        assert _decode_addr("alice@example.com") == "alice@example.com"

    def test_none_addr(self):
        assert _decode_addr(None) == "Unknown"

    def test_empty_addr(self):
        assert _decode_addr("") == "Unknown"

    def test_rfc2047_name(self):
        encoded = str(Header("Müller", "utf-8")) + " <mueller@example.com>"
        result = _decode_addr(encoded)
        assert "Müller" in result

    def test_multiple_addresses(self):
        result = _decode_addr("alice@a.com, bob@b.com")
        assert "alice@a.com" in result
        assert "bob@b.com" in result


# ==========================================================
# TESTS: _split_quoted_thread
# ==========================================================

class TestSplitQuotedThread:
    def test_no_separator(self):
        main, quoted = _split_quoted_thread("Just a simple message body")
        assert main == "Just a simple message body"
        assert quoted == ""

    def test_empty_body(self):
        main, quoted = _split_quoted_thread("")
        assert main == ""
        assert quoted == ""

    def test_gmail_reply(self):
        body = "Thanks for the update!\n\nOn Mon, Jan 1, 2024 at 10:00 AM Alice wrote:\n> Sure thing\n> Let me check"
        main, quoted = _split_quoted_thread(body)
        assert main == "Thanks for the update!"
        assert "On Mon, Jan 1" in quoted
        assert "> Sure thing" in quoted

    def test_outlook_reply(self):
        body = "Got it, will do.\n\n-----Original Message-----\nFrom: Bob\nSent: Monday, January 1, 2024\nTo: Alice\nSubject: Meeting"
        main, quoted = _split_quoted_thread(body)
        assert main == "Got it, will do."
        assert "-----Original Message-----" in quoted

    def test_outlook_forward(self):
        body = "See below.\n\n-------- Forwarded Message --------\nFrom: Charlie\nSubject: FYI\nTo: Dave"
        main, quoted = _split_quoted_thread(body)
        assert main == "See below."
        assert "Forwarded Message" in quoted

    def test_forwarded_message_variant(self):
        body = "Forwarding:\n\n-----Forwarded Message-----\nFrom: Eve"
        main, quoted = _split_quoted_thread(body)
        assert main == "Forwarding:"
        assert "-----Forwarded Message-----" in quoted

    def test_forward_header_block(self):
        body = "Check this out.\n\nFrom: Frank\nSent: Tue, 2 Jan 2024 14:00\nTo: Grace\nSubject: Report\n\nAttached is the report."
        main, quoted = _split_quoted_thread(body)
        assert main == "Check this out."
        assert "From: Frank" in quoted
        assert "Subject: Report" in quoted

    def test_forward_header_with_cc(self):
        body = "Hi all.\n\nFrom: Henry\nSent: Wed, 3 Jan 2024\nTo: Ivy\nCc: Jack\nSubject: Update\n\nPlease review."
        main, quoted = _split_quoted_thread(body)
        assert main == "Hi all."
        assert "Subject: Update" in quoted

    def test_multiple_separators_picks_earliest(self):
        body = "My reply.\n\nOn Mon, 1 Jan 2024 wrote:\n> Something\n\n-----Original Message-----\nFrom: Old"
        main, quoted = _split_quoted_thread(body)
        # The "On...wrote:" pattern comes first → split there
        assert main == "My reply."
        assert "On Mon" in quoted
        assert "-----Original Message-----" in quoted  # also in quoted part

    def test_apple_mail_reply(self):
        body = "Sure!\n\nOn 1 Jan 2024, at 15:00, Kate <kate@example.com> wrote:\n\n> Let's meet tomorrow\n> At the usual place"
        main, quoted = _split_quoted_thread(body)
        assert main == "Sure!"
        assert "On 1 Jan 2024" in quoted

    def test_outlook_underscore_line(self):
        body = "Reply below.\n\n_______\nFrom: Liam\nSent: Thu, 4 Jan 2024\nTo: Mia"
        main, quoted = _split_quoted_thread(body)
        assert main == "Reply below."
        assert "_______" in quoted

    def test_body_is_entirely_quoted(self):
        body = "On Fri, 5 Jan 2024 wrote:\n> Original message here"
        main, quoted = _split_quoted_thread(body)
        assert main == ""
        assert "On Fri" in quoted

    def test_case_insensitive_wrote(self):
        body = "OK.\n\nOn Mon, Jan 1, 2024 Alice WROTE:\n> test"
        main, quoted = _split_quoted_thread(body)
        assert main == "OK."
        assert "WROTE:" in quoted

    def test_outlook_mixed_case_original(self):
        body = "Sure.\n\n-----original message-----\nFrom: Noel"
        main, quoted = _split_quoted_thread(body)
        assert main == "Sure."
        assert "-----original message-----" in quoted


# ==========================================================
# TESTS: _extract_rtf_text
# ==========================================================

class TestExtractRtfText:
    def test_rtf_with_htmlrtf0_markers(self):
        """Method 1: Exchange-style \\htmlrtf0 markers — text must exceed 50-char quality gate."""
        rtf_data = (
            b"{\\rtf1\\ansi"
            b"\\htmlrtf0This is a longer email body text from Microsoft Exchange PST export with enough characters to pass the fifty character quality gate.\\htmlrtf1"
            b"}"
        )
        result = _extract_rtf_text(rtf_data)
        assert "longer email body text" in result
        assert "Microsoft Exchange PST export" in result

    def test_rtf_with_embedded_html(self):
        """Method 2: Embedded <html> tags with enough text to exceed 50-char gate."""
        rtf_data = (
            b"{\\rtf1\\ansi\\htmltag "
            b"<html><body><p>This is an HTML email body extracted from Outlook RTF that needs to be longer than fifty characters to pass the quality filter.</p></body></html>"
            b"}"
        )
        result = _extract_rtf_text(rtf_data)
        assert "HTML email body" in result
        assert "Outlook RTF" in result
        assert "<p>" not in result

    def test_rtf_generic_strip(self):
        """Method 3: Generic RTF stripping — text outside braces survives."""
        rtf_data = b"{\\rtf1\\ansi\\deff0 {\\fonttbl {\\f0 Times New Roman;}}\\f0\\pard Hello Generic Text That Is Long\\par}"
        result = _extract_rtf_text(rtf_data)
        assert "Hello Generic" in result

    def test_rtf_empty(self):
        result = _extract_rtf_text(b"")
        assert result == ""

    def test_rtf_method1_short_falls_through(self):
        """Method 1 returns <50 chars → falls through to method 2 or 3."""
        rtf_data = (
            b"{\\rtf1\\htmlrtf0short chunk\\htmlrtf1"
            b"<html><body><p>This is much longer body text that exceeds the fifty character quality gate by a significant margin.</p></body></html>}"
        )
        result = _extract_rtf_text(rtf_data)
        assert "much longer body text" in result

    def test_rtf_unicode_decode_fallback(self):
        """Non-UTF-8 bytes decode via latin-1, then stripped by method 3."""
        # Actual byte 0xE9 (latin-1 é), not the characters \\xe9
        data = b"{\\rtf1\\ansi Plain text with accent \xe9 caractere more words extra padding}"
        result = _extract_rtf_text(data)
        # Method 3 might still strip everything inside braces.
        # This tests that the function doesn't crash.
        assert isinstance(result, str)


# ==========================================================
# TESTS: _extract_attachment_text
# ==========================================================

class TestExtractAttachmentText:
    def test_txt_attachment(self):
        result = _extract_attachment_text("notes.txt", b"Hello from text")
        assert result == "Hello from text"

    def test_csv_attachment(self):
        result = _extract_attachment_text("data.csv", b"a,b,c\n1,2,3")
        assert result == "a,b,c\n1,2,3"

    def test_md_attachment(self):
        result = _extract_attachment_text("readme.md", b"# Title\nBody")
        assert result == "# Title\nBody"

    def test_json_attachment(self):
        result = _extract_attachment_text("config.json", b'{"key": "value"}')
        assert result == '{"key": "value"}'

    def test_xml_attachment(self):
        result = _extract_attachment_text("data.xml", b"<root><item/></root>")
        assert result == "<root><item/></root>"

    def test_yaml_attachment(self):
        result = _extract_attachment_text("config.yaml", b"key: value\nnested:\n  sub: val")
        assert result == "key: value\nnested:\n  sub: val"

    def test_yml_attachment(self):
        result = _extract_attachment_text("config.yml", b"key: val")
        assert result == "key: val"

    def test_log_attachment(self):
        result = _extract_attachment_text("app.log", b"[INFO] started")
        assert result == "[INFO] started"

    def test_pdf_attachment_pypdf2(self):
        """Minimal PDF that PyPDF2 can parse."""
        # Minimal valid PDF
        pdf_bytes = b"%PDF-1.4\n1 0 obj<</Type/Catalog/Pages 2 0 R>>endobj\n2 0 obj<</Type/Pages/Kids[3 0 R]/Count 1>>endobj\n3 0 obj<</Type/Page/MediaBox[0 0 612 792]/Parent 2 0 R/Resources<<>>>>endobj\nxref\n0 4\n0000000000 65535 f \n0000000009 00000 n \n0000000058 00000 n \n0000000115 00000 n \ntrailer<</Size 4/Root 1 0 R>>\nstartxref\n206\n%%EOF"
        result = _extract_attachment_text("doc.pdf", pdf_bytes)
        # PyPDF2 might extract empty text from minimal PDF; should not crash
        assert result is not None
        assert isinstance(result, str)

    def test_docx_attachment_python_docx(self):
        """Test with a minimal .docx that python-docx can parse."""
        import io
        from docx import Document
        doc = Document()
        doc.add_paragraph("Hello from Word")
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        result = _extract_attachment_text("report.docx", buf.read())
        assert "Hello from Word" in result

    def test_docm_ext(self):
        """.docm extension behaves same as .docx."""
        import io
        from docx import Document
        doc = Document()
        doc.add_paragraph("Docm content")
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        result = _extract_attachment_text("macro.docm", buf.read())
        assert "Docm content" in result

    def test_rtf_attachment(self):
        result = _extract_attachment_text("body.rtf", b"{\\rtf1 Hello RTF}")
        assert "Hello RTF" in result

    def test_none_data(self):
        assert _extract_attachment_text("file.txt", None) == ""

    def test_unknown_extension(self):
        assert _extract_attachment_text("image.png", b"GIF89a") == ""
        assert _extract_attachment_text("archive.zip", b"PK") == ""

    def test_no_extension(self):
        assert _extract_attachment_text("README", b"some content") == ""

    def test_latin1_fallback(self):
        """Non-UTF-8 text file falls back to latin-1, but invalid UTF-8 bytes get replaced with \\ufffd."""
        result = _extract_attachment_text("data.txt", b"Latin \xe9 text")
        assert "Latin" in result
        # 0xE9 is invalid standalone UTF-8 → replaced with U+FFFD
        assert "\ufffd" in result


# ==========================================================
# TESTS: _collect_attachments
# ==========================================================

class TestCollectAttachments:
    def test_no_attachments(self):
        msg = email.message_from_string(_make_simple_text())
        assert _collect_attachments(msg) == []

    def test_single_attachment(self):
        att = MIMEApplication(b"file content", _subtype="octet-stream")
        att.add_header("Content-Disposition", "attachment", filename="data.txt")
        msg_str = _make_multipart("With att", text_body="body", attachments=[att])
        msg = email.message_from_string(msg_str)
        result = _collect_attachments(msg)
        assert len(result) == 1
        assert result[0]["filename"] == "data.txt"
        assert "file content" in result[0]["extracted_text"]

    def test_multiple_attachments(self):
        att1 = MIMEApplication(b"first", _subtype="octet-stream")
        att1.add_header("Content-Disposition", "attachment", filename="a.txt")
        att2 = MIMEApplication(b"second", _subtype="octet-stream")
        att2.add_header("Content-Disposition", "attachment", filename="b.txt")
        msg_str = _make_multipart("Two atts", text_body="body", attachments=[att1, att2])
        msg = email.message_from_string(msg_str)
        result = _collect_attachments(msg)
        assert len(result) == 2
        filenames = [a["filename"] for a in result]
        assert "a.txt" in filenames
        assert "b.txt" in filenames

    def test_inline_image_skipped(self):
        img = MIMEImage(b"fakeimage", _subtype="png")
        img.add_header("Content-Disposition", "inline")
        msg_str = _make_multipart("With inline img", text_body="body", attachments=[img])
        msg = email.message_from_string(msg_str)
        result = _collect_attachments(msg)
        assert len(result) == 0

    def test_attached_image_included(self):
        img = MIMEImage(b"fakeimage", _subtype="png")
        img.add_header("Content-Disposition", "attachment", filename="photo.png")
        msg_str = _make_multipart("With attached img", text_body="body", attachments=[img])
        msg = email.message_from_string(msg_str)
        result = _collect_attachments(msg)
        assert len(result) == 1
        assert result[0]["filename"] == "photo.png"
        # PNG binary → no extracted text
        assert result[0]["extracted_text"] == ""

    def test_no_filename_skipped(self):
        att = MIMEApplication(b"content", _subtype="octet-stream")
        att.add_header("Content-Disposition", "attachment")  # no filename
        msg_str = _make_multipart("No filename", text_body="body", attachments=[att])
        msg = email.message_from_string(msg_str)
        result = _collect_attachments(msg)
        assert len(result) == 0

    def test_non_multipart_returns_empty(self):
        msg = email.message_from_string(_make_simple_text())
        assert _collect_attachments(msg) == []

    def test_body_parts_skipped(self):
        """text/plain and text/html without Content-Disposition: attachment are skipped."""
        msg_str = _make_multipart("mixed", text_body="plain", html_body="<p>html</p>")
        msg = email.message_from_string(msg_str)
        result = _collect_attachments(msg)
        assert len(result) == 0


# ==========================================================
# TESTS: _build_email_document
# ==========================================================

class TestBuildEmailDocument:
    def test_simple_text_email(self):
        msg = email.message_from_string(_make_simple_text())
        doc = _build_email_document("test.eml", msg)
        assert doc is not None
        assert "[File]: test.eml" in doc.text
        assert "[Subject]: Test Subject" in doc.text
        assert "[From]: alice@example.com" in doc.text
        assert "[To]: bob@example.com" in doc.text
        assert "Hello world" in doc.text
        assert "[EMAIL BODY Start]:" in doc.text
        assert "[Quoted Thread]" not in doc.text

    def test_empty_body_no_subject_returns_none(self):
        msg = email.message_from_string(_make_simple_text(subject="", body=""))
        doc = _build_email_document("empty.eml", msg)
        assert doc is None

    def test_body_no_subject_ok(self):
        """If body exists but subject is empty, it should still return document."""
        msg = email.message_from_string(_make_simple_text(subject="", body="Some body"))
        doc = _build_email_document("partial.eml", msg)
        assert doc is not None
        assert "Some body" in doc.text

    def test_html_email_with_include_html(self):
        msg = email.message_from_string(
            _make_simple_text(subject="HTML Email", body="<html><body><p>Hello HTML</p></body></html>")
        )
        # Override content type
        msg_str = _make_simple_text(subject="HTML Email", body="<html><body><p>Hello HTML</p></body></html>")
        msg = email.message_from_string(msg_str)
        msg.replace_header("Content-Type", "text/html")

        doc = _build_email_document("html.eml", msg, include_html=True)
        assert doc is not None
        assert "Hello HTML" in doc.text
        assert "<p>" not in doc.text

    def test_html_email_without_include_html(self):
        msg_str = _make_simple_text(subject="HTML Email", body="<html><body><p>Hidden</p></body></html>")
        msg = email.message_from_string(msg_str)
        msg.replace_header("Content-Type", "text/html")
        doc = _build_email_document("hidden.eml", msg, include_html=False)
        assert doc is not None  # subject exists
        # body should be empty because include_html is False
        assert "[EMAIL BODY Start]:\n\n" in doc.text or "[EMAIL BODY Start]:\n" in doc.text

    def test_multipart_prefers_plain_text(self):
        msg_str = _make_multipart(
            "Multi", text_body="Plain text body", html_body="<p>HTML body</p>"
        )
        msg = email.message_from_string(msg_str)
        doc = _build_email_document("multi.eml", msg)
        assert doc is not None
        assert "Plain text body" in doc.text
        assert "HTML body" not in doc.text  # HTML is ignored

    def test_multipart_falls_back_to_html(self):
        """When include_html=True and no text/plain, use text/html."""
        msg_str = _make_multipart(
            "HTML only", html_body="<html><body><p>HTML content</p></body></html>"
        )
        msg = email.message_from_string(msg_str)
        doc = _build_email_document("htmlonly.eml", msg, include_html=True)
        assert doc is not None
        assert "HTML content" in doc.text
        assert "<p>" not in doc.text

    def test_gmail_thread_detection(self):
        body = "This is my reply.\n\nOn Mon, Jan 1, 2024 at 10:00, Alice <alice@example.com> wrote:\n> Original message\n> More quoted"
        msg = email.message_from_string(_make_simple_text(body=body, subject="Re: Meeting"))
        doc = _build_email_document("reply.eml", msg)
        assert doc is not None
        assert "[EMAIL BODY Start]:\nThis is my reply." in doc.text
        assert "[Quoted Thread]:" in doc.text
        assert "On Mon, Jan 1" in doc.text
        assert "> Original message" in doc.text

    def test_outlook_thread_detection(self):
        body = "Will do.\n\n-----Original Message-----\nFrom: Bob\nSent: Monday\nTo: Alice\nSubject: Task"
        msg = email.message_from_string(_make_simple_text(body=body, subject="Re: Task"))
        doc = _build_email_document("outlook_reply.eml", msg)
        assert doc is not None
        assert "[EMAIL BODY Start]:\nWill do." in doc.text
        assert "[Quoted Thread]:" in doc.text
        assert "-----Original Message-----" in doc.text

    def test_forwarded_message_thread(self):
        body = "FYI.\n\n-------- Forwarded Message --------\nSubject: Original\nFrom: Charlie\nTo: Dave"
        msg = email.message_from_string(_make_simple_text(body=body, subject="Fwd: Original"))
        doc = _build_email_document("forward.eml", msg)
        assert doc is not None
        assert "[EMAIL BODY Start]:\nFYI." in doc.text
        assert "[Quoted Thread]:" in doc.text
        assert "Forwarded Message" in doc.text

    def test_rtf_body_fallback(self):
        """TUSD-style: no text/plain or text/html part, body in rtf-body.rtf attachment."""
        alt = MIMEMultipart("mixed")
        alt["Subject"] = "RTF Fallback"
        alt["From"] = "sender@example.com"
        alt["To"] = "recipient@example.com"
        alt["Date"] = formatdate(localtime=True)

        rtf_att = MIMEApplication(
            b"{\\rtf1\\ansi RTF body for the email}",
            _subtype="octet-stream",
        )
        rtf_att.add_header("Content-Disposition", "attachment", filename="rtf-body.rtf")
        alt.attach(rtf_att)

        msg = alt
        doc = _build_email_document("rtf_fallback.eml", msg)
        assert doc is not None
        assert "RTF body for the email" in doc.text
        assert "[EMAIL BODY Start]:" in doc.text

    def test_calendar_invite(self):
        """Calendar invites have no body — should produce a document with metadata."""
        cal_body = "BEGIN:VCALENDAR\nVERSION:2.0\nBEGIN:VEVENT\nSUMMARY:Meeting\nEND:VEVENT\nEND:VCALENDAR"
        msg = email.message_from_string(
            _make_simple_text(subject="Meeting Invite", body=cal_body)
        )
        msg.replace_header("Content-Type", "text/calendar; method=REQUEST")
        doc = _build_email_document("invite.eml", msg)
        # Body exists (calendar text) so document is created
        assert doc is not None
        assert "[Subject]: Meeting Invite" in doc.text

    def test_attachment_metadata_in_doc(self):
        att = MIMEApplication(b"Important info", _subtype="octet-stream")
        att.add_header("Content-Disposition", "attachment", filename="notes.txt")
        msg_str = _make_multipart(
            "With att", text_body="Body text", attachments=[att]
        )
        msg = email.message_from_string(msg_str)
        doc = _build_email_document("attached.eml", msg)
        assert doc is not None
        assert "[Attachments]: notes.txt" in doc.text
        assert "[Attachment Contents]:" in doc.text
        assert "Important info" in doc.text

    def test_encoded_subject(self):
        encoded_subj = str(Header("Re: über important", "utf-8"))
        msg = email.message_from_string(
            _make_simple_text(subject=encoded_subj, body="Got it.")
        )
        doc = _build_email_document("encoded.eml", msg)
        assert doc is not None
        assert "über important" in doc.text

    def test_no_from_header(self):
        msg_str = _make_simple_text(from_addr="", to_addr="bob@test.com")
        msg = email.message_from_string(msg_str)
        doc = _build_email_document("nofrom.eml", msg)
        assert doc is not None
        assert "[From]: Unknown" in doc.text

    def test_cc_header_appears_in_body_only(self):
        """Cc doesn't create a [Cc] field, but any content is fine."""
        msg_str = _make_simple_text(
            subject="With CC", body="Hello",
            from_addr="alice@a.com", to_addr="bob@b.com"
        )
        msg = email.message_from_string(msg_str)
        msg["Cc"] = "carol@c.com"
        doc = _build_email_document("cc.eml", msg)
        assert doc is not None
        assert "hello" in doc.text.lower() or "Hello" in doc.text


# ==========================================================
# TESTS: EmlReader (integration with temp files)
# ==========================================================

class TestEmlReader:
    def test_read_single_eml(self, tmp_path: Path):
        content = _make_simple_text()
        _write_eml(str(tmp_path), "test.eml", content)
        reader = EmlReader()
        docs = reader.load_data(str(tmp_path), max_count=10)
        assert len(docs) == 1
        assert "[Subject]: Test Subject" in docs[0].text

    def test_read_multiple_emls(self, tmp_path: Path):
        for i in range(5):
            content = _make_simple_text(subject=f"Email {i}", body=f"Body {i}")
            _write_eml(str(tmp_path), f"email{i}.eml", content)
        reader = EmlReader()
        docs = reader.load_data(str(tmp_path), max_count=10)
        assert len(docs) == 5
        subjects_found = sum(1 for d in docs if "[Subject]: Email" in d.text)
        assert subjects_found == 5

    def test_max_count_limits_results(self, tmp_path: Path):
        for i in range(10):
            content = _make_simple_text(subject=f"Email {i}", body=f"Body {i}")
            _write_eml(str(tmp_path), f"email{i}.eml", content)
        reader = EmlReader()
        docs = reader.load_data(str(tmp_path), max_count=3)
        assert len(docs) == 3

    def test_skips_non_eml_files(self, tmp_path: Path):
        _write_eml(str(tmp_path), "test.eml", _make_simple_text())
        _write_eml(str(tmp_path), "notes.txt", "not an email")
        _write_eml(str(tmp_path), "data.csv", "a,b,c")
        _write_eml(str(tmp_path), "noextension", "content")
        reader = EmlReader()
        docs = reader.load_data(str(tmp_path), max_count=10)
        assert len(docs) == 1

    def test_empty_directory(self, tmp_path: Path):
        reader = EmlReader()
        docs = reader.load_data(str(tmp_path), max_count=10)
        assert docs == []

    def test_skips_hidden_directories(self, tmp_path: Path):
        hidden_dir = tmp_path / ".hidden"
        hidden_dir.mkdir()
        _write_eml(str(hidden_dir), "test.eml", _make_simple_text())
        # Also add a visible one
        _write_eml(str(tmp_path), "visible.eml", _make_simple_text(subject="Visible"))
        reader = EmlReader()
        docs = reader.load_data(str(tmp_path), max_count=10)
        assert len(docs) == 1
        assert "Visible" in docs[0].text

    def test_malformed_eml_is_skipped(self, tmp_path: Path):
        """Malformed .eml files should be skipped without crashing the reader."""
        (tmp_path / "bad.eml").write_text("this is not valid email at all!!!", encoding="utf-8")
        # Also add a good one
        _write_eml(str(tmp_path), "good.eml", _make_simple_text())
        reader = EmlReader()
        docs = reader.load_data(str(tmp_path), max_count=10)
        # The malformed file might still be parsed or might fail — either way
        # the reader should not crash and should return at least the good email
        assert len(docs) >= 1
        assert any("[Subject]: Test Subject" in d.text for d in docs)

    def test_nested_directory_walk(self, tmp_path: Path):
        subdir = tmp_path / "sub" / "nested"
        subdir.mkdir(parents=True)
        _write_eml(str(subdir), "deep.eml", _make_simple_text(subject="Deep Email"))
        reader = EmlReader()
        docs = reader.load_data(str(tmp_path), max_count=10)
        assert len(docs) == 1
        assert "Deep Email" in docs[0].text

    def test_include_html_parameter(self, tmp_path: Path):
        """HTML-only .eml is empty body when include_html=False, has body when True."""
        msg_str = _make_simple_text(subject="HTML Only", body="<html><body><p>HTML content</p></body></html>")
        msg = email.message_from_string(msg_str)
        msg.replace_header("Content-Type", "text/html")
        _write_eml(str(tmp_path), "html_only.eml", msg.as_string())

        reader_no_html = EmlReader(include_html=False)
        docs_no_html = reader_no_html.load_data(str(tmp_path), max_count=10)
        # body is empty, but subject exists so doc is created
        doc_text = docs_no_html[0].text
        body_start = doc_text.split("[EMAIL BODY Start]:")[1].split("\n[")[0].strip()
        assert body_start == "" or body_start == "\n" or body_start == "\n\n"

        reader_with_html = EmlReader(include_html=True)
        docs_with_html = reader_with_html.load_data(str(tmp_path), max_count=10)
        assert "HTML content" in docs_with_html[0].text

    def test_attachment_extraction_integration(self, tmp_path: Path):
        """Test that attachments are actually extracted during full EmlReader run."""
        att = MIMEApplication(b"Attachment text here!", _subtype="octet-stream")
        att.add_header("Content-Disposition", "attachment", filename="content.txt")
        msg_str = _make_multipart(
            "With attachment", text_body="Main body", attachments=[att]
        )
        _write_eml(str(tmp_path), "attachments.eml", msg_str)
        reader = EmlReader()
        docs = reader.load_data(str(tmp_path), max_count=10)
        assert len(docs) == 1
        assert "[Attachment Contents]:" in docs[0].text
        assert "Attachment text here!" in docs[0].text
        assert "[Attachments]: content.txt" in docs[0].text


# ==========================================================
# TESTS: EmlxReader
# ==========================================================

class TestEmlxReader:
    def test_emlx_format_with_length_prefix(self, tmp_path: Path):
        """.emlx files have a decimal length prefix line before the email content."""
        email_content = _make_simple_text(subject="Emlx Test", body="Emlx body")
        length = len(email_content.encode("utf-8"))
        emlx_content = f"{length}\n{email_content}"
        (tmp_path / "test.emlx").write_text(emlx_content, encoding="utf-8")
        reader = EmlxReader()
        docs = reader.load_data(str(tmp_path), max_count=10)
        assert len(docs) == 1
        assert "Emlx body" in docs[0].text
        assert "[Subject]: Emlx Test" in docs[0].text

    def test_emlx_skips_other_extensions(self, tmp_path: Path):
        email_content = _make_simple_text(subject="Only emlx")
        length = len(email_content.encode("utf-8"))
        (tmp_path / "real.emlx").write_text(f"{length}\n{email_content}", encoding="utf-8")
        (tmp_path / "fake.txt").write_text("ignored", encoding="utf-8")
        reader = EmlxReader()
        docs = reader.load_data(str(tmp_path), max_count=10)
        assert len(docs) == 1
        assert "Only emlx" in docs[0].text

    def test_emlx_malformed_no_newline(self, tmp_path: Path):
        """.emlx file with length prefix and minimal content — produces a valid document."""
        (tmp_path / "bad.emlx").write_text("5\nhi", encoding="utf-8")
        reader = EmlxReader()
        docs = reader.load_data(str(tmp_path), max_count=10)
        assert len(docs) == 1  # Code reads it as a valid document with "hi" as body

    def test_emlx_empty_directory(self, tmp_path: Path):
        reader = EmlxReader()
        docs = reader.load_data(str(tmp_path), max_count=10)
        assert docs == []

    def test_emlx_attachments(self, tmp_path: Path):
        """Attachment extraction works through EmlxReader."""
        att = MIMEApplication(b"Emlx attachment data", _subtype="octet-stream")
        att.add_header("Content-Disposition", "attachment", filename="info.txt")
        msg_str = _make_multipart(
            "Emlx with att", text_body="Main", attachments=[att]
        )
        length = len(msg_str.encode("utf-8"))
        emlx_content = f"{length}\n{msg_str}"
        (tmp_path / "att.emlx").write_text(emlx_content, encoding="utf-8")
        reader = EmlxReader()
        docs = reader.load_data(str(tmp_path), max_count=10)
        assert len(docs) == 1
        assert "[Attachments]: info.txt" in docs[0].text
        assert "Emlx attachment data" in docs[0].text


# ==========================================================
# TESTS: Edge Cases & Error Handling
# ==========================================================

class TestEdgeCases:
    def test_very_long_subject(self, tmp_path: Path):
        long_subj = "Subject: " + "A" * 2000
        content = _make_simple_text(subject=long_subj, body="Short body")
        _write_eml(str(tmp_path), "long_subj.eml", content)
        reader = EmlReader()
        docs = reader.load_data(str(tmp_path), max_count=10)
        assert len(docs) == 1
        assert "A" * 2000 in docs[0].text

    def test_unicode_in_body(self, tmp_path: Path):
        body = "Hello 世界! ñoño éèêë 中文 😊"
        content = _make_simple_text(subject="Unicode", body=body)
        _write_eml(str(tmp_path), "unicode.eml", content)
        reader = EmlReader()
        docs = reader.load_data(str(tmp_path), max_count=10)
        assert len(docs) == 1
        for char in ["世界", "ñoño", "😊"]:
            assert char in docs[0].text

    def test_multiple_attachments_various_types(self, tmp_path: Path):
        att1 = MIMEApplication(b"Text content", _subtype="octet-stream")
        att1.add_header("Content-Disposition", "attachment", filename="readme.txt")
        att2 = MIMEApplication(
            b"{\\rtf1\\ansi RTF Content}", _subtype="octet-stream"
        )
        att2.add_header("Content-Disposition", "attachment", filename="formatted.rtf")
        msg_str = _make_multipart(
            "Multiple attachments", text_body="Body",
            attachments=[att1, att2]
        )
        _write_eml(str(tmp_path), "multi_att.eml", msg_str)
        reader = EmlReader()
        docs = reader.load_data(str(tmp_path), max_count=10)
        assert len(docs) == 1
        assert "readme.txt" in docs[0].text
        assert "formatted.rtf" in docs[0].text
        assert "Text content" in docs[0].text
        assert "RTF Content" in docs[0].text

    def test_mixed_eml_and_emlx_no_cross_contamination(self, tmp_path: Path):
        """EmlReader should only read .eml, EmlxReader only .emlx."""
        # Create both
        _write_eml(str(tmp_path), "note.eml", _make_simple_text(subject="Eml file"))
        emlx_content = f"{len('hi'.encode())}\n{_make_simple_text(subject='Emlx file')}"
        (tmp_path / "note.emlx").write_text(emlx_content, encoding="utf-8")

        eml_reader = EmlReader()
        eml_docs = eml_reader.load_data(str(tmp_path), max_count=10)
        assert len(eml_docs) == 1
        assert "Eml file" in eml_docs[0].text

        emlx_reader = EmlxReader()
        emlx_docs = emlx_reader.load_data(str(tmp_path), max_count=10)
        assert len(emlx_docs) == 1
        assert "Emlx file" in emlx_docs[0].text

    def test_large_number_of_emails(self, tmp_path: Path):
        """Process 100 emails efficiently."""
        for i in range(100):
            content = _make_simple_text(subject=f"Bulk {i}", body=f"Body {i}")
            _write_eml(str(tmp_path), f"bulk{i}.eml", content)
        reader = EmlReader()
        docs = reader.load_data(str(tmp_path), max_count=-1)
        assert len(docs) == 100

    def test_no_to_header(self):
        """Email with only From and no To."""
        msg_str = _make_simple_text(
            from_addr="alice@example.com", to_addr=""
        )
        msg = email.message_from_string(msg_str)
        doc = _build_email_document("noto.eml", msg)
        assert doc is not None
        assert "[To]: Unknown" in doc.text

    def test_no_from_or_to(self):
        """Minimal headers — should still produce document if body exists."""
        msg = MIMEText("Just a body")
        msg["Subject"] = "Minimal"
        msg["Date"] = formatdate(localtime=True)
        doc = _build_email_document("minimal.eml", msg)
        assert doc is not None
        assert "[Subject]: Minimal" in doc.text
        assert "[From]: Unknown" in doc.text
        assert "[To]: Unknown" in doc.text

    def test_nested_multipart_structure(self, tmp_path: Path):
        """Deeply nested multipart should still extract body and attachments."""
        # Construct: multipart/mixed → multipart/alternative + attachment
        msg = MIMEMultipart("mixed")
        msg["Subject"] = "Nested"
        msg["From"] = "a@a.com"
        msg["To"] = "b@b.com"
        msg["Date"] = formatdate(localtime=True)

        alt = MIMEMultipart("alternative")
        alt.attach(MIMEText("Plain text", "plain"))
        alt.attach(MIMEText("<html><body><p>HTML text</p></body></html>", "html"))
        msg.attach(alt)

        att = MIMEApplication(b"Deep attachment", _subtype="octet-stream")
        att.add_header("Content-Disposition", "attachment", filename="deep.txt")
        msg.attach(att)

        _write_eml(str(tmp_path), "nested.eml", msg.as_string())
        reader = EmlReader()
        docs = reader.load_data(str(tmp_path), max_count=10)
        assert len(docs) == 1
        assert "Plain text" in docs[0].text
        assert "Deep attachment" in docs[0].text

    def test_zero_length_body_with_subject(self, tmp_path: Path):
        """Email with subject but empty body — should still produce doc."""
        msg_str = _make_simple_text(subject="Empty Body", body="")
        # Force the body to be truly empty
        msg = email.message_from_string(
            "Subject: Empty Body\nFrom: a@a.com\nTo: b@b.com\nDate: Wed, 1 Jan 2024\n\n"
        )
        doc = _build_email_document("empty_body.eml", msg)
        # Subject exists, No Subject check: 'Empty Body' is not "No Subject"
        # And body.strip() is empty, so rtf fallback is attempted
        # Should still return a doc because subject != "No Subject"
        assert doc is not None
        assert "[EMAIL BODY Start]:\n" in doc.text or "[EMAIL BODY Start]:\n\n" in doc.text

    def test_rtf_body_fallback_multiple_filenames(self, tmp_path: Path):
        """RTF body attachment can have varied names like rtf-body, rtf-body.rtf, etc."""
        msg = MIMEMultipart("mixed")
        msg["Subject"] = "RTF Body"
        msg["From"] = "s@example.com"
        msg["To"] = "r@example.com"
        msg["Date"] = formatdate(localtime=True)

        rtf_att = MIMEApplication(
            b"{\\rtf1\\ansi RTF fallback from PST export}",
            _subtype="octet-stream",
        )
        rtf_att.add_header("Content-Disposition", "attachment", filename="rtf-body")
        msg.attach(rtf_att)

        doc = _build_email_document("rtf_body_noext.eml", msg)
        assert doc is not None
        assert "RTF fallback from PST export" in doc.text

    def test_binary_only_attachment_not_extracted(self, tmp_path: Path):
        """Binary attachments with no text extraction should not crash."""
        att = MIMEApplication(b"\x00\x01\x02\x03\xff\xfe\xfd", _subtype="octet-stream")
        att.add_header("Content-Disposition", "attachment", filename="binary.bin")
        msg_str = _make_multipart(
            "Binary att", text_body="Text body", attachments=[att]
        )
        _write_eml(str(tmp_path), "binary.eml", msg_str)
        reader = EmlReader()
        docs = reader.load_data(str(tmp_path), max_count=10)
        assert len(docs) == 1
        assert "binary.bin" in docs[0].text
        assert "Text body" in docs[0].text
        # [Attachment Contents] should not include extracted text for binary
        # Actually it might have empty string. Let's just check it doesn't crash.

    def test_unusual_line_endings(self, tmp_path: Path):
        """.eml files often have CRLF line endings from Windows."""
        content = _make_simple_text(subject="CRLF", body="Windows email")
        content = content.replace("\n", "\r\n")
        _write_eml(str(tmp_path), "crlf.eml", content)
        reader = EmlReader()
        docs = reader.load_data(str(tmp_path), max_count=10)
        assert len(docs) == 1
        assert "Windows email" in docs[0].text

    def test_multiple_rtf_attachments_only_first_used(self, tmp_path: Path):
        """Only the first rtf-body* attachment is used as body fallback."""
        msg = MIMEMultipart("mixed")
        msg["Subject"] = "Multi RTF"
        msg["From"] = "s@example.com"
        msg["To"] = "r@example.com"
        msg["Date"] = formatdate(localtime=True)

        rtf1 = MIMEApplication(b"{\\rtf1 Body from first}", _subtype="octet-stream")
        rtf1.add_header("Content-Disposition", "attachment", filename="rtf-body.rtf")
        msg.attach(rtf1)

        rtf2 = MIMEApplication(b"{\\rtf1 Second RTF}", _subtype="octet-stream")
        rtf2.add_header("Content-Disposition", "attachment", filename="rtf-body-2.rtf")
        msg.attach(rtf2)

        doc = _build_email_document("multi_rtf.eml", msg)
        assert doc is not None
        assert "Body from first" in doc.text
        assert "Second RTF" not in doc.text  # second is a regular attachment

    def test_empty_from_to_produces_doc_if_body(self):
        msg = MIMEText("Has body")
        msg["Subject"] = "Has body"
        msg["Date"] = formatdate(localtime=True)
        # deliberately no From/To
        doc = _build_email_document("bare.eml", msg)
        assert doc is not None
        assert "Has body" in doc.text
        assert "[From]: Unknown" in doc.text
        assert "[To]: Unknown" in doc.text

"""Email readers for LEANN — Apple Mail .emlx and standard .eml (RFC 2822).

Features:
- Metadata extraction (From, To, Subject, Date, attachments)
- HTML-to-plain-text via BeautifulSoup (when available)
- Reply/forward thread detection — main content prioritized over quoted thread
- Attachment content extraction (PDF, DOCX, TXT, RTF, CSV)
"""

import email
import os
import re
from email.header import decode_header
from pathlib import Path
from typing import Any

from llama_index.core import Document
from llama_index.core.readers.base import BaseReader

# ── Optional dependency imports ──────────────────────────
_bs4 = None
_PyPDF2 = None
_docx = None
try:
    from bs4 import BeautifulSoup as _bs4  # type: ignore
except ImportError:
    pass
try:
    import PyPDF2 as _PyPDF2  # type: ignore
except ImportError:
    try:
        import pypdf as _PyPDF2  # type: ignore
    except ImportError:
        pass
try:
    import docx as _docx  # type: ignore
except ImportError:
    pass


# ── Helpers ──────────────────────────────────────────────

def find_all_messages_directories(root: str | None = None) -> list[Path]:
    """Recursively find all 'Messages' directories under the given root."""
    if root is None:
        root = os.path.join(os.path.expanduser("~"), "Library", "Mail")
    return [
        Path(dp) for dp, _, _ in os.walk(root) if os.path.basename(dp) == "Messages"
    ]


def _payload_to_text(payload: object) -> str:
    if isinstance(payload, bytes):
        return payload.decode("utf-8", errors="ignore")
    if isinstance(payload, str):
        return payload
    return ""


def _strip_html(text: str) -> str:
    """Remove HTML tags and decode entities, returning plain text."""
    if _bs4 is not None:
        soup = _bs4(text, "html.parser")
        return soup.get_text(separator="\n")
    # Fallback: crude regex HTML stripping
    text = re.sub(r"<[^>]+>", "", text)
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def _decode_subject(subject: str | None) -> str:
    """Decode RFC 2047 encoded-word subjects to plain text."""
    if not subject:
        return "No Subject"
    try:
        parts = decode_header(subject)
        result = []
        for part, charset in parts:
            if isinstance(part, bytes):
                try:
                    result.append(part.decode(charset or "utf-8", errors="replace"))
                except (LookupError, UnicodeDecodeError):
                    result.append(part.decode("utf-8", errors="replace"))
            else:
                result.append(part)
        return " ".join(result)
    except Exception:
        return subject


def _decode_addr(addr: str | None) -> str:
    """Safely decode an address header."""
    if not addr:
        return "Unknown"
    try:
        parts = decode_header(addr)
        return "".join(
            p.decode(charset or "utf-8", errors="replace") if isinstance(p, bytes) else p
            for p, charset in parts
        )
    except Exception:
        return addr


# ── Reply/forward thread detection ──────────────────────

_REPLY_SEPARATORS = (
    # Gmail / Apple Mail:  On Mon, 1 Jan 2024 at 10:00, Person wrote:
    re.compile(r"^On\s+.+?\bwrote:\s*$", re.MULTILINE | re.IGNORECASE),
    # Outlook:  -----Original Message-----
    re.compile(r"^-{3,}Original\s+Message-{3,}\s*$", re.MULTILINE | re.IGNORECASE),
    # Outlook underscore line (3+ underscores)
    re.compile(r"^_{5,}\s*$", re.MULTILINE),
    # Forward header block:  From: ...  Sent: ...  To: ...  Subject: ...
    re.compile(
        r"^From:.+\nSent:.+\nTo:.+\n(?:Cc:.+\n)?(?:Subject:.+\n)",
        re.MULTILINE | re.IGNORECASE,
    ),
    # Forwarded message:  ----- Forwarded Message -----  or similar
    re.compile(
        r"^-{3,}\s*Forward(?:ed)?\s+Message\s*-{3,}\s*$",
        re.MULTILINE | re.IGNORECASE,
    ),
)


def _split_quoted_thread(body: str) -> tuple[str, str]:
    """Split email body into (main_content, quoted_thread).

    Returns (body, "") if no reply/forward separator is found.
    The quoted text is everything *from* the first separator onward.
    """
    best_pos = len(body)
    for pat in _REPLY_SEPARATORS:
        m = pat.search(body)
        if m and m.start() < best_pos:
            best_pos = m.start()

    if best_pos < len(body):
        main_content = body[:best_pos].strip()
        quoted = body[best_pos:].strip()
        return main_content, quoted

    return body.strip(), ""


# ── Attachment text extraction ──────────────────────────

def _extract_rtf_text(data: bytes) -> str:
    """Extract text from RTF data.

    Strips RTF control words, braces, and MS Exchange/Outlook-specific
    markers (\htmlrtf, \htmltag, \bkmkstart, etc.) to get readable text.
    """
    try:
        text = data.decode("utf-8", errors="replace")
    except Exception:
        text = data.decode("latin-1", errors="replace")

    # Method 1: Extract text from \htmlrtf0 markers (Microsoft Exchange)
    # These mark the actual email text within the RTF.
    pieces = re.findall(r"\\htmlrtf0\\s*([^\\}]+)", text)
    if pieces:
        result = " ".join(p.strip() for p in pieces)
        # Strip any remaining RTF artifacts
        result = re.sub(r"[{}]", " ", result)
        result = re.sub(r"<[^>]+>", " ", result)
        result = re.sub(r"\s+", " ", result).strip()
        if len(result) > 50:
            return result

    # Method 2: Try to find embedded HTML (Outlook Exchange RTF with <html> tags)
    html_match = re.search(r"<html[^>]*>.*?</html>", text, re.DOTALL | re.IGNORECASE)
    if html_match:
        raw_html = html_match.group()
        if _bs4 is not None:
            soup = _bs4(raw_html, "html.parser")
            result = soup.get_text(separator="\n").strip()
        else:
            result = re.sub(r"<[^>]+>", " ", raw_html)
            result = re.sub(r"\s+", " ", result).strip()
        # Clean up common RTF artifacts that leak through
        result = re.sub(r"[{}]", " ", result)
        result = re.sub(r"\\[a-z*]+\d*\b", " ", result)
        result = re.sub(r"\s+", " ", result).strip()
        if len(result) > 50:
            return result

    # Method 3: Strip all RTF markup generically
    groups = re.sub(r"\{[^}]*\}", " ", text)
    cleaned = re.sub(r"\\(?:[a-z]+|[*]|'[0-9a-f]{2})-?\d*\b", " ", groups)
    cleaned = re.sub(r"[{}]", " ", cleaned)
    cleaned = re.sub(r"<[^>]+>", " ", cleaned)
    cleaned = re.sub(r"\s+", " ", cleaned).strip()
    return cleaned

def _extract_attachment_text(
    filename: str, data: bytes | None
) -> str:
    """Extract text content from a supported attachment type. Returns '' on failure."""
    if not data:
        return ""

    ext = Path(filename).suffix.lower()

    # Plain text
    if ext in (".txt", ".csv", ".log", ".md", ".json", ".xml", ".yaml", ".yml"):
        try:
            return data.decode("utf-8", errors="replace")
        except Exception:
            return data.decode("latin-1", errors="replace")

    # RTF
    if ext == ".rtf":
        return _extract_rtf_text(data)

    # PDF
    if ext == ".pdf":
        if _PyPDF2 is not None:
            try:
                import io

                reader = _PyPDF2.PdfReader(io.BytesIO(data))
                return "\n".join(
                    page.extract_text() or "" for page in reader.pages
                )
            except Exception:
                return ""
        return ""

    # Word
    if ext in (".docx", ".docm"):
        if _docx is not None:
            try:
                import io

                doc = _docx.Document(io.BytesIO(data))
                return "\n".join(p.text for p in doc.paragraphs)
            except Exception:
                return ""
        return ""

    return ""


def _collect_attachments(msg: Any) -> list[dict]:
    """Walk MIME parts and return list of {filename, data, extracted_text}.

    Skips inline images and the email body parts (text/plain, text/html).
    """
    attachments: list[dict] = []
    if not msg.is_multipart():
        return attachments

    for part in msg.walk():
        ctype = part.get_content_type()
        cdispo = str(part.get("Content-Disposition", ""))

        # Skip body parts
        if ctype in ("text/plain", "text/html") and "attachment" not in cdispo:
            continue

        # Skip inline images
        if ctype.startswith("image/") and "attachment" not in cdispo:
            continue

        filename = part.get_filename()
        if not filename:
            continue

        raw = part.get_payload(decode=True)
        text = _extract_attachment_text(filename, raw)

        attachments.append({
            "filename": filename,
            "size": len(raw) if raw else 0,
            "content_type": ctype,
            "extracted_text": text,
        })

    return attachments


# ── Document builder ─────────────────────────────────────

def _build_email_document(
    filename: str,
    msg: Any,
    include_html: bool = False,
    attachments_dir: str | None = None,
) -> Document | None:
    """Build a LEANN Document from an email.message.Message.

    Handles:
    - Subject decoding (RFC 2047)
    - HTML-to-plain-text conversion
    - Reply/forward thread separation
    - Attachment metadata + text extraction
    """
    subject = _decode_subject(msg.get("Subject"))
    from_addr = _decode_addr(msg.get("From"))
    to_addr = _decode_addr(msg.get("To"))
    date = msg.get("Date", "Unknown")

    # ── Extract body ──────────────────────────────────────
    # Try text/plain first; fall back to text/html if include_html is set
    body = ""
    if msg.is_multipart():
        for part in msg.walk():
            ctype = part.get_content_type()
            if ctype == "text/plain":
                payload = part.get_payload(decode=True)
                if payload:
                    body = _payload_to_text(payload)
                    break
        # Fall back to HTML if no plain text found
        if not body and include_html:
            for part in msg.walk():
                ctype = part.get_content_type()
                if ctype == "text/html":
                    payload = part.get_payload(decode=True)
                    if payload:
                        raw = _payload_to_text(payload)
                        body = _strip_html(raw)
                        break
    else:
        payload = msg.get_payload(decode=True)
        if payload:
            raw = _payload_to_text(payload)
            ctype = msg.get_content_type()
            if ctype == "text/html" or "html" in (raw[:50].lower() if raw else ""):
                body = _strip_html(raw) if include_html else ""
            else:
                body = raw

    if not body.strip() and subject == "No Subject":
        return None

    # ── RTF body fallback ──────────────────────────────────
    # Some PST exports (e.g. TUSD's) have no text/plain or
    # text/html part — the HTML body was converted to an
    # rtf-body.rtf attachment. Use it as the email body.
    attachments = []
    if not body.strip():
        all_attachments = _collect_attachments(msg)
        for a in all_attachments:
            if a["filename"].lower().startswith("rtf-body") and a["extracted_text"]:
                body = a["extracted_text"]
                break
        if not body.strip():
            attachments = all_attachments
    else:
        attachments = _collect_attachments(msg)

    # ── Thread separation ────────────────────────────────
    main_body, quoted_thread = _split_quoted_thread(body)

    # ── Build document text ───────────────────────────────
    doc_text = (
        f"\n[File]: {filename}\n"
        f"[From]: {from_addr}\n"
        f"[To]: {to_addr}\n"
        f"[Subject]: {subject}\n"
        f"[Date]: {date}\n"
    )

    if attachments:
        att_names = "; ".join(a["filename"] for a in attachments)
        doc_text += f"[Attachments]: {att_names}\n"

    doc_text += f"[EMAIL BODY Start]:\n{main_body}\n"

    if quoted_thread:
        doc_text += f"\n[Quoted Thread]:\n{quoted_thread}\n"

    if attachments:
        doc_text += "\n[Attachment Contents]:\n"
        for a in attachments:
            if a["extracted_text"]:
                doc_text += (
                    f"--- {a['filename']} ---\n"
                    f"{a['extracted_text'][:5000]}\n\n"
                )

    return Document(text=doc_text, metadata={})


# ── Readers ──────────────────────────────────────────────

class EmlxReader(BaseReader):
    """Apple Mail .emlx file reader.

    .emlx files have a length prefix line followed by RFC 2822 content.
    """

    def __init__(
        self,
        include_html: bool = False,
        attachments_dir: str | None = None,
    ) -> None:
        self.include_html = include_html
        self.attachments_dir = attachments_dir

    def load_data(self, input_dir: str, **load_kwargs: Any) -> list[Document]:
        docs: list[Document] = []
        max_count = load_kwargs.get("max_count", 1000)
        count = total_files = successful_files = failed_files = 0

        print(f"Starting to process directory: {input_dir}")

        for dirpath, dirnames, filenames in os.walk(input_dir):
            dirnames[:] = [d for d in dirnames if not d.startswith(".")]
            for filename in filenames:
                if max_count > 0 and count >= max_count:
                    break
                if not filename.endswith(".emlx"):
                    continue

                total_files += 1
                filepath = os.path.join(dirpath, filename)
                try:
                    with open(filepath, encoding="utf-8", errors="ignore") as f:
                        content = f.read()

                    lines = content.split("\n", 1)
                    if len(lines) < 2:
                        continue

                    msg = email.message_from_string(lines[1])
                    doc = _build_email_document(
                        filename, msg, self.include_html, self.attachments_dir
                    )
                    if doc:
                        docs.append(doc)
                        count += 1
                        successful_files += 1
                        if successful_files <= 3:
                            subj = re.sub(r"\s+", " ", doc.text.split("\n[Subject]: ", 1)[-1].split("\n")[0] if "[Subject]:" in doc.text else "?")[:50]
                            print(f"Loaded: {filename} - Subject: {subj}...")

                except Exception as e:
                    failed_files += 1
                    if failed_files <= 5:
                        print(f"Error parsing {filepath}: {e}")
                    continue

        print(f"  Total .emlx: {total_files}  Loaded: {successful_files}  Failed: {failed_files}")
        return docs


class EmlReader(BaseReader):
    """Standard .eml (RFC 2822) email file reader.

    Unlike .emlx, .eml files have no length prefix — the entire file is the email.
    """

    def __init__(
        self,
        include_html: bool = False,
        attachments_dir: str | None = None,
    ) -> None:
        self.include_html = include_html
        self.attachments_dir = attachments_dir

    def load_data(self, input_dir: str, **load_kwargs: Any) -> list[Document]:
        docs: list[Document] = []
        max_count = load_kwargs.get("max_count", 1000)
        count = total_files = successful_files = failed_files = 0

        print(f"Starting to process directory: {input_dir}")

        for dirpath, dirnames, filenames in os.walk(input_dir):
            dirnames[:] = [d for d in dirnames if not d.startswith(".")]
            for filename in filenames:
                if max_count > 0 and count >= max_count:
                    break
                if not filename.endswith(".eml"):
                    continue

                total_files += 1
                filepath = os.path.join(dirpath, filename)
                try:
                    with open(filepath, encoding="utf-8", errors="ignore") as f:
                        content = f.read()

                    msg = email.message_from_string(content)
                    doc = _build_email_document(
                        filename, msg, self.include_html, self.attachments_dir
                    )
                    if doc:
                        docs.append(doc)
                        count += 1
                        successful_files += 1
                        if successful_files <= 3:
                            subj = re.sub(r"\s+", " ", doc.text.split("\n[Subject]: ", 1)[-1].split("\n")[0] if "[Subject]:" in doc.text else "?")[:50]
                            print(f"Loaded: {filename} - Subject: {subj}...")

                except Exception as e:
                    failed_files += 1
                    if failed_files <= 5:
                        print(f"Error parsing {filepath}: {e}")
                    continue

        print(f"  Total .eml: {total_files}  Loaded: {successful_files}  Failed: {failed_files}")
        return docs
